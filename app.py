from flask import Flask, render_template, request, send_file
import requests
from bs4 import BeautifulSoup
from io import BytesIO
from utils.despacho import gerar_despacho
from docx import Document
import os

app = Flask(__name__)

NORMAS_URLS = [
    {"nome": "Lei 8.212/91", "url": "https://www.planalto.gov.br/ccivil_03/leis/l8212cons.htm"},
    {"nome": "Lei 8.213/91", "url": "https://www.planalto.gov.br/ccivil_03/leis/l8213cons.htm"},
    {"nome": "Lei 8.742/93 (LOAS)", "url": "https://www.planalto.gov.br/ccivil_03/leis/l8742.htm"},
    {"nome": "Decreto 3.048/99", "url": "https://www.planalto.gov.br/ccivil_03/decreto/d3048.htm"},
    {"nome": "Decreto 6.214/07", "url": "https://www.planalto.gov.br/ccivil_03/_ato2007-2010/2007/decreto/d6214.htm"},
    {"nome": "Lei 10.779/03", "url": "https://www.planalto.gov.br/ccivil_03/leis/2003/L10.779compilado.htm"},
    {"nome": "Portal INSS", "url": "https://portalin.inss.gov.br/"},
]

def busca_termo(url, termo, artigo=None):
    resultados = []
    try:
        resp = requests.get(url, timeout=15)
        resp.encoding = resp.apparent_encoding
        soup = BeautifulSoup(resp.text, "html.parser")
        texto = soup.get_text(separator=" ", strip=True)
        linhas = texto.split(". ")
        for linha in linhas:
            if termo.lower() in linha.lower():
                if artigo:
                    if artigo.lower() in linha.lower():
                        resultados.append(linha.strip())
                else:
                    resultados.append(linha.strip())
    except Exception as e:
        resultados.append(f"Erro ao acessar {url}: {str(e)}")
    return resultados

@app.route("/", methods=["GET", "POST"])
def index():
    resultados = []
    termo = ""
    artigo = ""
    norma_escolhida = ""
    despacho_texto = ""
    tipo_despacho = ""
    exportar = False

    if request.method == "POST":
        termo = request.form.get("termo", "").strip()
        artigo = request.form.get("artigo", "").strip()
        norma_escolhida = request.form.get("norma", "")
        tipo_despacho = request.form.get("tipo_despacho", "")
        exportar = request.form.get("exportar", "") == "on"

        urls_para_busca = NORMAS_URLS
        if norma_escolhida:
            urls_para_busca = [n for n in NORMAS_URLS if n["nome"] == norma_escolhida]

        for norma in urls_para_busca:
            trechos = busca_termo(norma["url"], termo, artigo)
            if trechos:
                resultados.append({"nome": norma["nome"], "url": norma["url"], "trechos": trechos})

        if tipo_despacho:
            despacho_texto = gerar_despacho(tipo_despacho, termo, artigo, resultados)

        if exportar:
            doc = Document()
            doc.add_heading('Resultados da Consulta', level=1)
            doc.add_paragraph(f'Termo pesquisado: {termo}')
            if artigo:
                doc.add_paragraph(f'Artigo/Palavra-chave: {artigo}')
            for resultado in resultados:
                doc.add_heading(resultado["nome"], level=2)
                for trecho in resultado["trechos"]:
                    doc.add_paragraph(trecho, style='List Bullet')
            if despacho_texto:
                doc.add_heading('Despacho Gerado', level=1)
                doc.add_paragraph(despacho_texto)
            fp = BytesIO()
            doc.save(fp)
            fp.seek(0)
            return send_file(fp, as_attachment=True, download_name="resultado_normas.docx")
    return render_template(
        "index.html",
        resultados=resultados,
        termo=termo,
        artigo=artigo,
        norma_escolhida=norma_escolhida,
        despacho_texto=despacho_texto,
        tipo_despacho=tipo_despacho
    )

if __name__ == "__main__":
    os.makedirs("utils", exist_ok=True)
    app.run(debug=True)
